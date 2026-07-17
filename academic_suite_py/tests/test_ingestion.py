import pytest
import io
import fitz
import docx
import asyncio
from fastapi.testclient import TestClient
from app.main import app
from app.services.extraction import (
    extract_document,
    extract_pasted_text,
    validate_input_file,
    IngestionError
)
from app.services.cleaning import clean_and_normalize_text
from app.services.hashing import generate_document_hash
from app.db import (
    initialize_database,
    create_scan,
    get_scan_by_hash,
    get_scan_by_id
)

client = TestClient(app)

@pytest.fixture(scope="module", autouse=True)
def setup_db():
    # Initialize the sqlite database before running tests
    loop = asyncio.get_event_loop()
    loop.run_until_complete(initialize_database())

# Helper to generate programmatic PDF bytes
def create_test_pdf(text: str) -> bytes:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), text)
    pdf_bytes = doc.write()
    doc.close()
    return pdf_bytes

# Helper to generate programmatic DOCX bytes
def create_test_docx(text: str, cell_text: str = "") -> bytes:
    doc = docx.Document()
    doc.add_paragraph(text)
    if cell_text:
        table = doc.add_table(rows=1, cols=1)
        table.rows[0].cells[0].text = cell_text
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def test_pasted_text_extraction():
    res = extract_pasted_text("This is a simple pasted academic text snippet.")
    assert res["source_type"] == "pasted_text"
    assert res["character_count"] > 0
    assert res["word_count"] == 8

def test_empty_pasted_text():
    with pytest.raises(IngestionError) as exc_info:
        extract_pasted_text("   ")
    assert exc_info.value.code == "EMPTY_FILE"

def test_txt_extraction():
    content = "Hello standard text".encode("utf-8")
    res = extract_document("test.txt", content)
    assert res["source_type"] == "txt"
    assert "Hello standard text" in res["raw_text"]

def test_utf8_bom_txt():
    content = "\ufeffHello UTF-8 BOM text".encode("utf-8")
    res = extract_document("test.txt", content)
    assert res["source_type"] == "txt"
    assert "Hello UTF-8 BOM text" in res["raw_text"]

def test_empty_txt():
    content = b"   "
    with pytest.raises(IngestionError) as exc_info:
        extract_document("test.txt", content)
    assert exc_info.value.code == "EMPTY_FILE"

def test_unsupported_extension():
    content = b"some data"
    with pytest.raises(IngestionError) as exc_info:
        extract_document("test.png", content)
    assert exc_info.value.code == "UNSUPPORTED_FILE_TYPE"

def test_invalid_mime_type():
    # PNG signature with txt extension
    content = b"\x89PNG\r\n\x1a\n"
    with pytest.raises(IngestionError) as exc_info:
        extract_document("test.txt", content)
    assert exc_info.value.code == "MALFORMED_DOCUMENT"

def test_oversized_upload():
    large_content = b"a" * (5 * 1024 * 1024 + 1)
    with pytest.raises(IngestionError) as exc_info:
        validate_input_file("test.txt", large_content)
    assert exc_info.value.code == "FILE_TOO_LARGE"

def test_valid_docx_extraction():
    content = create_test_docx("Ingestion DOCX paragraph.", "Table cell value.")
    res = extract_document("test.docx", content)
    assert res["source_type"] == "docx"
    assert "Ingestion DOCX paragraph" in res["raw_text"]
    assert "Table cell value" in res["raw_text"]

def test_valid_pdf_extraction():
    content = create_test_pdf("Ingestion PDF test string content.")
    res = extract_document("test.pdf", content)
    assert res["source_type"] == "pdf"
    assert "Ingestion PDF test string content" in res["raw_text"]
    assert res["page_count"] == 1

def test_scanned_pdf_detection():
    # PDF with empty/scanned content
    doc = fitz.open()
    doc.new_page()  # Blank page
    pdf_bytes = doc.write()
    doc.close()
    
    res = extract_document("scanned.pdf", pdf_bytes)
    assert "scanned_pdf_detected" in res["warnings"]
    assert "No extractable text" in res["raw_text"]

def test_text_cleaning():
    dirty = "  Unicode \u201cquotes\u201d\n\n\n\r\nand  multiple    spaces.  "
    cleaned = clean_and_normalize_text(dirty)
    # Check unicode quotes normalize to standard quotes or remain clean
    assert "quotes" in cleaned["cleaned_text"]
    # Multiple spaces collapsed
    assert "multiple spaces." in cleaned["cleaned_text"]
    # Gaps collapsed to exactly 2 newlines
    assert cleaned["cleaned_text"].count("\n\n") == 1
    assert not cleaned["cleaned_text"].startswith(" ")

def test_whitespace_normalization():
    dirty = "   line 1   \n   line 2   "
    cleaned = clean_and_normalize_text(dirty)
    assert cleaned["cleaned_text"] == "line 1\nline 2"

def test_citation_preservation():
    citation = "As shown in [1] and discussed by Smith et al. (2024), this is correct."
    cleaned = clean_and_normalize_text(citation)
    assert "[1]" in cleaned["cleaned_text"]
    assert "(2024)" in cleaned["cleaned_text"]

def test_sha256_consistency():
    t1 = "Normalized text 1."
    t2 = "Normalized text 1."
    assert generate_document_hash(t1) == generate_document_hash(t2)

def test_sha256_difference():
    t1 = "Normalized text 1."
    t2 = "Normalized text 2."
    assert generate_document_hash(t1) != generate_document_hash(t2)

@pytest.mark.asyncio
async def test_database_insert_and_lookup():
    doc_hash = generate_document_hash("Unique text for database insert test.")
    scan = await create_scan(
        user_id="test-user-123",
        filename="db_test.txt",
        file_type="txt",
        mime_type="text/plain",
        document_hash=doc_hash,
        original_text="Unique text for database insert test.",
        cleaned_text="Unique text for database insert test.",
        character_count=37,
        word_count=6
    )
    assert scan.id is not None
    
    fetched = await get_scan_by_hash(doc_hash)
    assert fetched is not None
    assert fetched.id == scan.id
    assert fetched.user_id == "test-user-123"

@pytest.mark.asyncio
async def test_cache_hit_on_repeated_document():
    doc_hash = generate_document_hash("Repeated document body for caching test.")
    # First insert
    await create_scan(
        user_id="test-user-123",
        filename="cache_test.txt",
        file_type="txt",
        mime_type="text/plain",
        document_hash=doc_hash,
        original_text="Repeated document body for caching test.",
        cleaned_text="Repeated document body for caching test.",
        character_count=40,
        word_count=6
    )
    
    # Endpoint test with mock auth headers
    headers = {"Authorization": "Bearer mock-user-123"}
    data = {"pasted_text": "Repeated document body for caching test."}
    
    response = client.post("/api/v1/documents/extract", data=data, headers=headers)
    assert response.status_code == 200
    json_data = response.json()
    assert json_data["cache_hit"] is True
    assert "scan_id" in json_data

def test_successful_extract_endpoint():
    headers = {"Authorization": "Bearer mock-user-123"}
    data = {"pasted_text": "This is a brand new pasted text that has never been scanned."}
    response = client.post("/api/v1/documents/extract", data=data, headers=headers)
    assert response.status_code == 200
    json_data = response.json()
    assert json_data["cache_hit"] is False
    assert "scan_id" in json_data
    assert json_data["document"]["word_count"] == 12

def test_invalid_extract_endpoint_request():
    headers = {"Authorization": "Bearer mock-user-123"}
    # Both provided
    data = {"pasted_text": "Some text", "file": ("test.txt", b"Some data")}
    response = client.post("/api/v1/documents/extract", data=data, headers=headers)
    assert response.status_code == 400
    assert "error" in response.json()

def test_safe_error_response():
    headers = {"Authorization": "Bearer mock-user-123"}
    # Unsupported extension
    files = {"file": ("test.png", b"Fake png")}
    response = client.post("/api/v1/documents/extract", files=files, headers=headers)
    assert response.status_code == 415
    json_data = response.json()
    assert "error" in json_data
    assert json_data["error"]["code"] == "UNSUPPORTED_FILE_TYPE"
    assert "png" not in json_data["error"]["message"].lower()  # Safe validation error
