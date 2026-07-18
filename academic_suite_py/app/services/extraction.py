import fitz  # PyMuPDF
import docx  # python-docx
import io
import os
from typing import Dict, Any, List, Optional

MAX_FILE_SIZE = 5 * 1024 * 1024  # 5 MB


class IngestionError(Exception):
    def __init__(self, code: str, message: str, status_code: int = 400):
        self.code = code
        self.message = message
        self.status_code = status_code
        super().__init__(message)


def validate_input_file(filename: Optional[str], content: bytes) -> str:
    """Validate file metadata and content signature.
    Returns the validated MIME type or raises IngestionError.
    """
    if not content or len(content) == 0:
        raise IngestionError("EMPTY_FILE", "The uploaded file is empty.", 400)

    if len(content) > MAX_FILE_SIZE:
        raise IngestionError(
            "FILE_TOO_LARGE",
            f"File exceeds the maximum size of {MAX_FILE_SIZE // (1024*1024)}MB.",
            413,
        )

    if not filename:
        raise IngestionError("VALIDATION_ERROR", "Filename is missing.", 400)

    _, ext = os.path.splitext(filename.lower())
    if ext not in {".pdf", ".docx", ".txt"}:
        raise IngestionError(
            "UNSUPPORTED_FILE_TYPE", "Only PDF, DOCX, and TXT files are supported.", 415
        )

    # Check magic numbers/signatures
    detected_mime = None
    if content.startswith(b"%PDF"):
        detected_mime = "application/pdf"
    elif content.startswith(b"PK\x03\x04"):
        # Zip signature (applies to DOCX)
        detected_mime = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        # Check if it can be decoded as plain text
        try:
            content.decode("utf-8")
            detected_mime = "text/plain"
        except UnicodeDecodeError:
            try:
                content.decode("utf-8-sig")
                detected_mime = "text/plain"
            except UnicodeDecodeError:
                pass

    if not detected_mime:
        raise IngestionError(
            "MALFORMED_DOCUMENT", "Could not verify document content signature.", 422
        )

    # Verify extension matches detected MIME
    if ext == ".pdf" and detected_mime != "application/pdf":
        raise IngestionError(
            "MALFORMED_DOCUMENT",
            "File extension (.pdf) does not match detected PDF content.",
            422,
        )
    if (
        ext == ".docx"
        and detected_mime
        != "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        raise IngestionError(
            "MALFORMED_DOCUMENT",
            "File extension (.docx) does not match detected DOCX content.",
            422,
        )
    if ext == ".txt" and detected_mime != "text/plain":
        raise IngestionError(
            "MALFORMED_DOCUMENT",
            "File extension (.txt) does not match plain text content.",
            422,
        )

    return detected_mime


def extract_pdf(content: bytes, filename: str) -> Dict[str, Any]:
    warnings: List[str] = []
    try:
        doc = fitz.open(stream=content, filetype="pdf")
    except Exception:
        raise IngestionError(
            "MALFORMED_DOCUMENT", "The PDF file is malformed or corrupt.", 422
        )

    pages_text = []
    page_count = len(doc)

    for page_num in range(page_count):
        try:
            page = doc.load_page(page_num)
            text = page.get_text()
            pages_text.append(text)
        except Exception:
            pages_text.append("")

    raw_text = "\n--- PAGE BOUNDARY ---\n".join(pages_text)

    # Detect scanned PDF: if no extractable text (raw_text stripped is empty)
    if page_count > 0 and not raw_text.strip():
        raw_text = "No extractable text was found. This document may be scanned. OCR is not enabled in the current PAGGY MVP."
        warnings.append("scanned_pdf_detected")

    char_count = len(raw_text)
    word_count = len(raw_text.split())

    return {
        "source_type": "pdf",
        "filename": filename,
        "mime_type": "application/pdf",
        "raw_text": raw_text,
        "page_count": page_count,
        "character_count": char_count,
        "word_count": word_count,
        "warnings": warnings,
    }


def extract_docx(content: bytes, filename: str) -> Dict[str, Any]:
    try:
        doc = docx.Document(io.BytesIO(content))
    except Exception:
        raise IngestionError(
            "MALFORMED_DOCUMENT", "The DOCX file is malformed or corrupt.", 422
        )

    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())

    raw_text = "\n\n".join(paragraphs)
    if not raw_text.strip():
        raise IngestionError(
            "NO_EXTRACTABLE_TEXT", "The DOCX file contains no extractable text.", 422
        )

    char_count = len(raw_text)
    word_count = len(raw_text.split())

    return {
        "source_type": "docx",
        "filename": filename,
        "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "raw_text": raw_text,
        "page_count": 0,
        "character_count": char_count,
        "word_count": word_count,
        "warnings": [],
    }


def extract_txt(content: bytes, filename: str) -> Dict[str, Any]:
    decodings = ["utf-8", "utf-8-sig", "latin1"]
    raw_text = None

    for encoding in decodings:
        try:
            raw_text = content.decode(encoding)
            break
        except UnicodeDecodeError:
            continue

    if raw_text is None:
        # Fallback with replacement characters
        raw_text = content.decode("utf-8", errors="replace")

    if not raw_text.strip():
        raise IngestionError("EMPTY_FILE", "The plain text file is empty.", 400)

    char_count = len(raw_text)
    word_count = len(raw_text.split())

    return {
        "source_type": "txt",
        "filename": filename,
        "mime_type": "text/plain",
        "raw_text": raw_text,
        "page_count": 0,
        "character_count": char_count,
        "word_count": word_count,
        "warnings": [],
    }


def extract_pasted_text(text: str) -> Dict[str, Any]:
    if not text or not text.strip():
        raise IngestionError("EMPTY_FILE", "Pasted text is empty.", 400)

    char_count = len(text)
    word_count = len(text.split())

    return {
        "source_type": "pasted_text",
        "filename": "pasted_text.txt",
        "mime_type": "text/plain",
        "raw_text": text,
        "page_count": 0,
        "character_count": char_count,
        "word_count": word_count,
        "warnings": [],
    }


def extract_document(filename: Optional[str], content: bytes) -> Dict[str, Any]:
    mime_type = validate_input_file(filename, content)
    if mime_type == "application/pdf":
        return extract_pdf(content, filename)
    elif (
        mime_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        return extract_docx(content, filename)
    else:
        return extract_txt(content, filename)
